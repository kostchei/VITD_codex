from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Wastes_Rules_Summary_and_Pseudocode.docx'

def set_font(run, size=11, bold=False, color='000000', name='Calibri'):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:ascii'), name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in('w:tblW')
    tblW.set(qn('w:w'), '9360'); tblW.set(qn('w:type'), 'dxa')
    grid = table._tbl.tblGrid
    for col, w in zip(grid.gridCol_lst, widths): col.set(qn('w:w'), str(w))
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Inches(w/1440)
            tcW = cell._tc.tcPr.tcW
            tcW.set(qn('w:w'), str(w)); tcW.set(qn('w:type'), 'dxa')

def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix); set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):]); set_font(r)
    else:
        r = p.add_run(text); set_font(r)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text))

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    p.paragraph_format.space_after = Pt(10 if level == 1 else 7)
    set_font(p.add_run(text), 16 if level == 1 else 13, True, '2E74B5')
    return p

def code_line(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25 * indent)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    set_font(p.add_run(text), 9, False, '1F2937', 'Consolas')
    return p

doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
sec.header_distance = sec.footer_distance = Inches(0.492)

# compact_reference_guide tokens
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Calibri'; normal._element.rPr.rFonts.set(qn('w:ascii'), 'Calibri'); normal._element.rPr.rFonts.set(qn('w:hAnsi'), 'Calibri')
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

# memo_masthead selected: restrained technical-reference title block
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run('RULES REFERENCE'), 10, True, '5B6573')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
set_font(p.add_run('Wastes / Decay Engine'), 24, True, '0B2545')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(16)
set_font(p.add_run('Summary and implementation pseudocode from the supplied game-rule sheet'), 11, False, '5B6573')

heading(doc, 'What this system does')
add_body(doc, 'The sheet defines a card-driven consequence engine for Wasteland-style travel. Players draw a Wastes card, resolve its numbered rule, then continue until the card or scenario says otherwise. The entries combine travel hazards, resource loss, wounds, random encounters, combat, and occasionally a beneficial discovery.')
add_body(doc, 'The rules are intentionally stateful: a result can change health, water, supplies, weapons, armour, companions, current location, and active conditions such as injury, illness, radiation, or an escorting NPC.')

heading(doc, 'Core rule patterns')
for item in [
    'Random resolution: many entries call for a die roll, sometimes with a modifier. The total selects a row on the card.',
    'Tests and saves: the sheet repeatedly asks for skill/attribute checks (for example Intelligence, Strength, Survival, Medicine, Luck, or a relevant combat skill). A pass avoids, reduces, or redirects the penalty; a failure applies the stated consequence.',
    'Resource pressure: hazards consume water, supplies, ammunition, armour integrity, or weapon condition. Some effects also force an item discard or an immediate return to a settlement.',
    'Health and conditions: outcomes deal wounds, inflict ongoing conditions, or reduce a character statistic. Conditions should be stored separately from immediate damage so later rules can reference them.',
    'Encounters: results may create an NPC, creature, raider, or settlement interaction. These may branch to dialogue, combat, trade, an escort obligation, or a new objective.',
    'Combat/escape: hostile results can start a fight. If a rule permits retreat or escape, it normally changes position and may still impose damage or loss.',
    'Special results: a few outcomes grant a discovery, treasure, information, a shortcut, temporary aid, or a story hook rather than a penalty.'
]: add_bullet(doc, item)

heading(doc, 'Resolution model', 1)
table = doc.add_table(rows=1, cols=2)
set_table_widths(table, [2700, 6660])
hdr = table.rows[0].cells
for c, text in zip(hdr, ['Phase', 'Engine responsibility']):
    shade(c, 'E8EEF5'); set_cell_margins(c); set_font(c.paragraphs[0].add_run(text), 10, True, '1F4D78')
rows = [
    ('1. Draw', 'Select or draw a Wastes card and identify its numbered entry.'),
    ('2. Parse', 'Identify required rolls, tests, choices, costs, targets, and follow-up instructions.'),
    ('3. Resolve', 'Apply results in order: roll -> modifier -> test/choice -> consequence.'),
    ('4. Update', 'Commit all changes to the party and world state, including ongoing conditions and spawned encounters.'),
    ('5. Continue', 'Start combat, move/return the party, add a hook, or end the encounter as instructed.')
]
for a,b in rows:
    cells = table.add_row().cells
    for c, text in zip(cells, [a,b]):
        set_cell_margins(c)
        set_font(c.paragraphs[0].add_run(text), 10, c == cells[0], '000000')

heading(doc, 'Pseudocode', 1)
add_body(doc, 'This pseudocode is designed for a digital adaptation. Individual card entries should be data-driven rule definitions rather than hard-coded branches wherever possible.')
for line, indent in [
    ('function resolveWastesEncounter(party, world):', 0),
    ('card = drawWastesCard(world.deck)', 1),
    ('entry = chooseOrRollEntry(card, party, world)', 1),
    ('log("Wastes: " + card.name + " / " + entry.title)', 1),
    ('', 0),
    ('context = { party, world, card, entry, pendingEffects: [] }', 1),
    ('for step in entry.steps in listed order:', 1),
    ('result = resolveStep(step, context)', 2),
    ('context.pendingEffects.append(result.effects)', 2),
    ('if result.endsEncounter:', 2),
    ('break', 3),
    ('', 0),
    ('applyEffectsAtomically(context.pendingEffects, party, world)', 1),
    ('runFollowUps(entry, party, world)  // combat, movement, NPC, objective, discard, etc.', 1),
    ('return buildEncounterReport(card, entry, context)', 1),
    ('', 0),
    ('function resolveStep(step, context):', 0),
    ('if step.type == "roll":', 1),
    ('total = rollDice(step.dice) + getModifier(step, context.party)', 2),
    ('return selectOutcome(step.outcomes, total)', 2),
    ('if step.type == "test":', 1),
    ('score = rollTest(step.skill, step.difficulty, context.party)', 2),
    ('return score >= step.difficulty ? step.onPass : step.onFail', 2),
    ('if step.type == "choice":', 1),
    ('choice = requestPlayerChoice(step.options)', 2),
    ('return step.outcomes[choice]', 2),
    ('if step.type == "combat":', 1),
    ('return startCombat(step.enemyGroup, context.party, context.world)', 2),
    ('return evaluateRule(step, context)', 1)
]:
    code_line(doc, line, indent)

heading(doc, 'Implementation notes', 1)
for item in [
    'Keep card text and outcome tables as content data. The engine should only know generic operations such as roll, test, apply damage, lose resource, add condition, spawn encounter, and move party.',
    'Validate costs before committing an effect. If a player cannot pay a stated cost, follow the card fallback; if none exists, flag the rule as needing a designer decision rather than silently going negative.',
    'Record a resolution log with the card, dice, modifiers, selected choice, and all state changes. This makes tabletop adjudication and debugging reproducible.',
    'Treat ambiguous wording in the source image as authoring input that should be verified against the original full-resolution rules document before shipping an automated ruleset.'
]: add_bullet(doc, item)

doc.save(OUT)
print(OUT)
